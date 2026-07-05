#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
從 charter.html 直接產生正式排版的 Word（.docx）與 PDF：
標題／章名／條號／內文全篇統一標楷體，用字級大小與深藍色區分層次（不切換字體）。

先決條件（Ubuntu/Debian，僅需裝一次）：
    apt-get install -y fonts-arphic-ukai
    pip install python-docx reportlab

用法：
    python3 tools/gen_charter_docs.py

輸出：
    高雄榮民總醫院企業工會章程（草案）.docx
    高雄榮民總醫院企業工會章程（草案）.pdf
（覆蓋 repo 根目錄的同名檔案；記得重新 bump.sh + commit + push）
"""
import re
import pathlib

REPO = pathlib.Path(__file__).resolve().parent.parent
CHARTER_HTML = REPO / "charter.html"
OUT_BASENAME = "高雄榮民總醫院企業工會章程（草案）"

UKAI_PATH = "/usr/share/fonts/truetype/arphic/ukai.ttc"
UKAI_SUBFONT_INDEX = 2  # "AR PL UKai TW"（標楷體風格）

TITLE_TEXT = "高雄榮民總醫院企業工會章程（草案）"
SUBTITLE_TEXT = "（草案版本，尚未經成立大會通過；提供籌備委員會及成立大會審議用）"
CLOSING_DATE_LINE = "本章程經中華民國【　】年【　】月【　】日成立大會通過。"
CLOSING_SIGN_LINE = "理事長：楊淯涵"

NAVY_HEX = "1A3A6B"


def extract_charter_body() -> str:
    text = CHARTER_HTML.read_text(encoding="utf-8")
    start = text.index('<h2>第一章　總則</h2>')
    end = text.index('<p>本章程經中華民國')
    return text[start:end]


def parse_tokens(body: str):
    """回傳 [(kind, a, b), ...]；kind='chapter' → a=章名；kind='article' → a=條號, b=內容(含<br>)"""
    tokens = []
    for h2, artnum, content in re.findall(
        r'<h2>(.*?)</h2>|<p class="art"><b>(.*?)</b>　(.*?)</p>', body, re.S
    ):
        if h2:
            tokens.append(("chapter", h2, None))
        else:
            tokens.append(("article", artnum, content))
    return tokens


def gen_pdf(tokens, out_path: pathlib.Path):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib import colors

    pdfmetrics.registerFont(TTFont("Kai", UKAI_PATH, subfontIndex=UKAI_SUBFONT_INDEX))

    NAVY = colors.HexColor(f"#{NAVY_HEX}")
    BLACK = colors.HexColor("#111111")
    GRAY = colors.HexColor("#666666")

    title_style = ParagraphStyle("title", fontName="Kai", fontSize=23, leading=33,
                                  alignment=TA_CENTER, spaceAfter=10, textColor=BLACK)
    subtitle_style = ParagraphStyle("subtitle", fontName="Kai", fontSize=10.5, leading=17,
                                     alignment=TA_CENTER, textColor=GRAY, spaceAfter=22)
    chapter_style = ParagraphStyle("chapter", fontName="Kai", fontSize=16, leading=23,
                                    spaceBefore=22, spaceAfter=12, textColor=NAVY, alignment=TA_CENTER)
    art_style = ParagraphStyle("art", fontName="Kai", fontSize=12.5, leading=22,
                                spaceBefore=7, spaceAfter=4, alignment=TA_JUSTIFY, textColor=BLACK)
    sub_style = ParagraphStyle("sub", fontName="Kai", fontSize=12.5, leading=22,
                                spaceAfter=4, leftIndent=22, alignment=TA_JUSTIFY, textColor=BLACK)
    plain_style = ParagraphStyle("plain", fontName="Kai", fontSize=12.5, leading=22, spaceAfter=4,
                                  alignment=TA_JUSTIFY)

    story = [
        Paragraph(TITLE_TEXT, title_style),
        Paragraph(SUBTITLE_TEXT, subtitle_style),
        HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#B9C4D6"), spaceAfter=18),
    ]

    for kind, a, b in tokens:
        if kind == "chapter":
            story.append(Paragraph(a, chapter_style))
            continue
        lines = b.split("<br>")
        first = lines[0]
        story.append(Paragraph(f'<font color="#{NAVY_HEX}">{a}</font>　{first}', art_style))
        for extra in lines[1:]:
            extra = extra.strip()
            if extra:
                story.append(Paragraph(extra, sub_style))

    story.append(Spacer(1, 16 * mm))
    story.append(Paragraph(CLOSING_DATE_LINE, plain_style))
    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph(CLOSING_SIGN_LINE, plain_style))

    doc = SimpleDocTemplate(str(out_path), pagesize=A4, topMargin=26 * mm, bottomMargin=23 * mm,
                             leftMargin=27 * mm, rightMargin=27 * mm)
    doc.build(story)


def gen_docx(tokens, out_path: pathlib.Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    KAI = "標楷體"
    NAVY = RGBColor(0x1A, 0x3A, 0x6B)
    GRAY = RGBColor(0x66, 0x66, 0x66)

    def set_font(run, size, color=None, font=KAI):
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = False
        if color:
            run.font.color.rgb = color
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font)

    def add_para(text, size=12.5, align=None, color=None, space_before=0, space_after=6,
                 line_spacing=1.35):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = line_spacing
        set_font(p.add_run(text), size, color)
        return p

    def add_art_para(artnum, content, size=12.5, space_before=5, space_after=4, line_spacing=1.35):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = line_spacing
        set_font(p.add_run(artnum), size, NAVY)
        set_font(p.add_run(f"　{content}"), size, None)
        return p

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Mm(25), Mm(22)
    sec.left_margin, sec.right_margin = Mm(28), Mm(28)

    add_para(TITLE_TEXT, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, line_spacing=1.2)
    add_para(SUBTITLE_TEXT, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY,
             space_after=20, line_spacing=1.2)

    for kind, a, b in tokens:
        if kind == "chapter":
            add_para(a, size=16, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=18, space_after=10, line_spacing=1.2)
            continue
        lines = b.split("<br>")
        first = lines[0]
        add_art_para(a, first)
        for extra in lines[1:]:
            extra = extra.strip()
            if extra:
                p = add_para(extra, size=12.5, space_after=4)
                p.paragraph_format.left_indent = Mm(6)

    add_para("", size=12.5, space_after=20)
    add_para(CLOSING_DATE_LINE, size=12.5, space_after=10)
    add_para(CLOSING_SIGN_LINE, size=12.5, space_after=4)

    doc.save(str(out_path))


def main():
    body = extract_charter_body()
    tokens = parse_tokens(body)
    print(f"parsed {len(tokens)} tokens (chapters+articles)")

    pdf_path = REPO / f"{OUT_BASENAME}.pdf"
    docx_path = REPO / f"{OUT_BASENAME}.docx"

    gen_pdf(tokens, pdf_path)
    print("wrote", pdf_path)
    gen_docx(tokens, docx_path)
    print("wrote", docx_path)


if __name__ == "__main__":
    main()
